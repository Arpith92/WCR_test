import os
import pandas as pd
from docxtpl import DocxTemplate
from datetime import datetime
import streamlit as st
import zipfile
import io
from pathlib import Path
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.styles import getSampleStyleSheet

# ---- Constants ----
TEMPLATE_DOC = "sample.docx"   # Template must be in repo
OUT_DIR = "Result"
os.makedirs(OUT_DIR, exist_ok=True)

def _safe(x):
    """Format values safely"""
    if pd.isna(x) or x == "":
        return ""
    if isinstance(x, (datetime, pd.Timestamp)):
        return x.strftime("%d-%m-%Y")
    try:
        num = float(x)
        return f"{num:.2f}"
    except (ValueError, TypeError):
        return str(x).strip()

# ----------------- Streamlit UI -----------------
st.title("📑 Automated WCR Generator & Converter")

# ----------------- Step 1 -----------------
st.subheader("Step 1: Generate Word Files from Excel")

uploaded_file = st.file_uploader("Upload Input Excel File", type=["xlsx"], key="excel_upload")

if uploaded_file is not None:
    df = pd.read_excel(uploaded_file, engine="openpyxl")
    df.columns = df.columns.str.strip()

    # Rename for consistency
    rename_map = {
        "wo no": "wo_no", "wo_no": "wo_no",
        "wo date": "wo_date", "wo_date": "wo_date",
        "wo des": "wo_des", "wo_des": "wo_des",
        "location_code": "Location_code", "Location_code": "Location_code",
        "customername_code": "customername_code",
        "capacity_code": "Capacity_code", "Capacity_code": "Capacity_code",
        "site_incharge": "site_incharge",
        "Scada_incharge": "Scada_incharge",
        "Re_date": "Re_date",
        "Site_Name": "Site_Name",
        "Line_1_Workstatus": "Line_1_Workstatus",
        "Line_2_Workstatus": "Line_2_Workstatus",
        "Payment Terms": "Payment_Terms"
    }
    df = df.rename(columns={c: rename_map.get(c, c) for c in df.columns})

    generated_word = []

    for i, row in df.iterrows():
        context = {col: _safe(row[col]) for col in df.columns}

        # --- Auto-generate Sr. No. ---
        for n in [1, 2, 3]:
            fields = [
                context.get(f"Line_{n}", ""),
                context.get(f"Line_{n}_WO_qty", ""),
                context.get(f"Line_{n}_UOM", ""),
                context.get(f"Line_{n}_PB_qty", ""),
                context.get(f"Line_{n}_TB_Qty", ""),
                context.get(f"Line_{n}_cu_qty", ""),
                context.get(f"Line_{n}_B_qty", "")
            ]
            context[f"item_sr_no_{n}"] = str(n) if any(f for f in fields) else ""

        # Render Word with docxtpl
        doc = DocxTemplate(TEMPLATE_DOC)
        doc.render(context)
        wo = context.get("wo_no", "") or f"Row{i+1}"
        word_path = os.path.join(OUT_DIR, f"WCR_{wo}.docx")
        doc.save(word_path)
        generated_word.append(word_path)

    # ---- ZIP Word ----
    zip_word = io.BytesIO()
    with zipfile.ZipFile(zip_word, "w") as zipf:
        for file in generated_word:
            zipf.write(file, arcname=os.path.basename(file))
    zip_word.seek(0)

    st.download_button(
        "⬇️ Download All WCR Files (Word ZIP)",
        data=zip_word,
        file_name="WCR_Word_Files.zip",
        mime="application/zip"
    )

# ----------------- Step 2 -----------------
st.subheader("Step 2: Convert Word ZIP to PDF ZIP")

uploaded_zip = st.file_uploader("Upload WCR Word ZIP", type=["zip"], key="word_zip_upload")

if uploaded_zip is not None:
    extract_dir = Path("TempExtract")
    extract_dir.mkdir(exist_ok=True)

    # Extract uploaded ZIP
    with zipfile.ZipFile(uploaded_zip, "r") as zip_ref:
        zip_ref.extractall(extract_dir)

    pdf_files = []
    for docx_file in extract_dir.glob("*.docx"):
        pdf_path = extract_dir / (docx_file.stem + ".pdf")

        # Read Word content
        doc = Document(docx_file)
        story = []
        styles = getSampleStyleSheet()
        story.append(Paragraph(f"Converted: {docx_file.name}", styles['Title']))
        story.append(Spacer(1, 12))

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text, styles['Normal']))
                story.append(Spacer(1, 6))

        # Tables
        for table in doc.tables:
            data = []
            for row in table.rows:
                data.append([cell.text.strip() for cell in row.cells])
            if data:
                story.append(Table(data))
                story.append(Spacer(1, 12))

        pdf = SimpleDocTemplate(str(pdf_path))
        pdf.build(story)

        pdf_files.append(pdf_path)

    if pdf_files:
        # Create ZIP of PDFs
        zip_pdf = io.BytesIO()
        with zipfile.ZipFile(zip_pdf, "w") as zipf:
            for file in pdf_files:
                zipf.write(file, arcname=os.path.basename(file))
        zip_pdf.seek(0)

        st.success(f"✅ Converted {len(pdf_files)} Word files to PDF (layout simplified)")
        st.download_button(
            "⬇️ Download All PDFs (ZIP)",
            data=zip_pdf,
            file_name="WCR_PDF_Files.zip",
            mime="application/zip"
        )
